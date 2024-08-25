from flask import Flask, render_template, request, send_file, abort
import io
from PIL import Image
import ffmpeg
import subprocess
import os
from docx import Document
from PyPDF2 import PdfReader, PdfWriter
import shutil

application = Flask(__name__)

@application.route('/', methods=['GET'])
def index():
    return render_template('index.html')

def convert_with_libreoffice(input_path, output_format):
    command = ['libreoffice', '--headless', '--convert-to', output_format, '--outdir', os.path.dirname(input_path), input_path]
    subprocess.run(command, check=True)

def convert_doc_to_pdf(input_path, output_path):
    # Use LibreOffice to convert .doc to .pdf
    command = ['libreoffice', '--headless', '--convert-to', 'pdf', '--outdir', os.path.dirname(output_path), input_path]
    subprocess.run(command, check=True)

# Image converters
@application.route('/jpgtopng', methods=['GET', 'POST'])
def jpg_to_png():
    if request.method == 'POST':
        try:
            image = request.files['file']
            img = Image.open(image)
            png_io = io.BytesIO()
            img.save(png_io, format='PNG')
            png_io.seek(0)
            return send_file(png_io, mimetype='image/png', as_attachment=True, download_name='converted.png')
        except Exception as e:
            print(f"Error: {e}")
            abort(500)
    return render_template('upload.html', convert='jpg_to_png')

@application.route('/pngtojpg', methods=['GET', 'POST'])
def png_to_jpg():
    if request.method == 'POST':
        try:
            image = request.files['file']
            img = Image.open(image)
            jpg_io = io.BytesIO()
            img.convert('RGB').save(jpg_io, format='JPEG')
            jpg_io.seek(0)
            return send_file(jpg_io, mimetype='image/jpeg', as_attachment=True, download_name='converted.jpg')
        except Exception as e:
            print(f"Error: {e}")
            abort(500)
    return render_template('upload.html', convert='png_to_jpg')

@application.route('/webptopng', methods=['GET', 'POST'])
def webp_to_png():
    if request.method == 'POST':
        try:
            image = request.files['file']
            img = Image.open(image)
            png_io = io.BytesIO()
            img.save(png_io, format='PNG')
            png_io.seek(0)
            return send_file(png_io, mimetype='image/png', as_attachment=True, download_name='converted.png')
        except Exception as e:
            print(f"Error: {e}")
            abort(500)
    return render_template('upload.html', convert='webp_to_png')

@application.route('/bmptopng', methods=['GET', 'POST'])
def bmp_to_png():
    if request.method == 'POST':
        try:
            image = request.files['file']
            img = Image.open(image)
            png_io = io.BytesIO()
            img.save(png_io, format='PNG')
            png_io.seek(0)
            return send_file(png_io, mimetype='image/png', as_attachment=True, download_name='converted.png')
        except Exception as e:
            print(f"Error: {e}")
            abort(500)
    return render_template('upload.html', convert='bmp_to_png')

@application.route('/pngtopdf', methods=['GET', 'POST'])
def png_to_pdf():
    if request.method == 'POST':
        try:
            image = request.files['file']
            img = Image.open(image)
            pdf_io = io.BytesIO()
            img.save(pdf_io, format='PDF')
            pdf_io.seek(0)
            return send_file(pdf_io, mimetype='application/pdf', as_attachment=True, download_name='converted.pdf')
        except Exception as e:
            print(f"Error: {e}")
            abort(500)
    return render_template('upload.html', convert='png_to_pdf')

@application.route('/jpgtopdf', methods=['GET', 'POST'])
def jpg_to_pdf():
    if request.method == 'POST':
        try:
            image = request.files['file']
            img = Image.open(image)
            pdf_io = io.BytesIO()
            img.save(pdf_io, format='PDF')
            pdf_io.seek(0)
            return send_file(pdf_io, mimetype='application/pdf', as_attachment=True, download_name='converted.pdf')
        except Exception as e:
            print(f"Error: {e}")
            abort(500)
    return render_template('upload.html', convert='jpg_to_pdf')


@application.route('/mp4tomp3', methods=['GET', 'POST'])
def mp4_to_mp3():
    if request.method == 'POST':
        try:
            audio_file = request.files['file']
            input_path = 'input.mp4'
            output_path = 'output.mp3'
            audio_file.save(input_path)
            ffmpeg.input(input_path).output(output_path).overwrite_output().run()
            return send_file(output_path, as_attachment=True, download_name='converted.mp3')
        except Exception as e:
            print(f"Error: {e}")
            abort(500)
    return render_template('upload.html', convert='mp4_to_mp3')

@application.route('/mp3toflac', methods=['GET', 'POST'])
def mp3_to_flac():
    if request.method == 'POST':
        try:
            audio_file = request.files['file']
            input_path = 'input.mp3'
            output_path = 'output.flac'
            audio_file.save(input_path)
            ffmpeg.input(input_path).output(output_path).run()
            return send_file(output_path, as_attachment=True, download_name='converted.flac')
        except Exception as e:
            print(f"Error: {e}")
            abort(500)
    return render_template('upload.html', convert='mp3_to_flac')

@application.route('/flactomp3', methods=['GET', 'POST'])
def flac_to_mp3():
    if request.method == 'POST':
        try:
            audio_file = request.files['file']
            input_path = 'input.flac'
            output_path = 'output.mp3'
            audio_file.save(input_path)
            ffmpeg.input(input_path).output(output_path).run()
            return send_file(output_path, as_attachment=True, download_name='converted.mp3')
        except Exception as e:
            print(f"Error: {e}")
            abort(500)
    return render_template('upload.html', convert='flac_to_mp3')

@application.route('/mp3toogg', methods=['GET', 'POST'])
def mp3_to_ogg():
    if request.method == 'POST':
        try:
            audio_file = request.files['file']
            input_path = 'input.mp3'
            output_path = 'output.ogg'
            audio_file.save(input_path)
            ffmpeg.input(input_path).output(output_path).run()
            return send_file(output_path, as_attachment=True, download_name='converted.ogg')
        except Exception as e:
            print(f"Error: {e}")
            abort(500)
    return render_template('upload.html', convert='mp3_to_ogg')

@application.route('/mp3towav', methods=['GET', 'POST'])
def mp3_to_wav():
    if request.method == 'POST':
        try:
            audio_file = request.files['file']
            input_path = 'input.mp3'
            output_path = 'output.wav'
            audio_file.save(input_path)
            ffmpeg.input(input_path).output(output_path).run()
            return send_file(output_path, as_attachment=True, download_name='converted.wav')
        except Exception as e:
            print(f"Error: {e}")
            abort(500)
    return render_template('upload.html', convert='mp3_to_wav')

# @application.route('/mp4tomp3', methods=['GET', 'POST'])
# def mp4_to_mp3():
#     if request.method == 'POST':
#         try:
#             audio_file = request.files['file']
#             input_path = '/tmp/input.mp4'
#             output_path = '/tmp/output.mp3'
#             audio_file.save(input_path)
#             ffmpeg.input(input_path).output(output_path).overwrite_output().run()
#             return send_file(output_path, as_attachment=True, download_name='converted.mp3')
#         except Exception as e:
#             print(f"Error: {e}")
#             abort(500)
#         finally:
#             if os.path.exists(input_path):
#                 os.remove(input_path)
#             if os.path.exists(output_path):
#                 os.remove(output_path)
#     return render_template('upload.html', convert='mp4_to_mp3')

# @application.route('/mp3toflac', methods=['GET', 'POST'])
# def mp3_to_flac():
#     if request.method == 'POST':
#         try:
#             audio_file = request.files['file']
#             input_path = '/tmp/input.mp3'
#             output_path = '/tmp/output.flac'
#             audio_file.save(input_path)
#             ffmpeg.input(input_path).output(output_path).run()
#             return send_file(output_path, as_attachment=True, download_name='converted.flac')
#         except Exception as e:
#             print(f"Error: {e}")
#             abort(500)
#         finally:
#             if os.path.exists(input_path):
#                 os.remove(input_path)
#             if os.path.exists(output_path):
#                 os.remove(output_path)
#     return render_template('upload.html', convert='mp3_to_flac')

# @application.route('/flactomp3', methods=['GET', 'POST'])
# def flac_to_mp3():
#     if request.method == 'POST':
#         try:
#             audio_file = request.files['file']
#             input_path = '/tmp/input.flac'
#             output_path = '/tmp/output.mp3'
#             audio_file.save(input_path)
#             ffmpeg.input(input_path).output(output_path).run()
#             return send_file(output_path, as_attachment=True, download_name='converted.mp3')
#         except Exception as e:
#             print(f"Error: {e}")
#             abort(500)
#         finally:
#             if os.path.exists(input_path):
#                 os.remove(input_path)
#             if os.path.exists(output_path):
#                 os.remove(output_path)
#     return render_template('upload.html', convert='flac_to_mp3')

# @application.route('/mp3toogg', methods=['GET', 'POST'])
# def mp3_to_ogg():
#     if request.method == 'POST':
#         try:
#             audio_file = request.files['file']
#             input_path = '/tmp/input.mp3'
#             output_path = '/tmp/output.ogg'
#             audio_file.save(input_path)
#             ffmpeg.input(input_path).output(output_path).run()
#             return send_file(output_path, as_attachment=True, download_name='converted.ogg')
#         except Exception as e:
#             print(f"Error: {e}")
#             abort(500)
#         finally:
#             if os.path.exists(input_path):
#                 os.remove(input_path)
#             if os.path.exists(output_path):
#                 os.remove(output_path)
#     return render_template('upload.html', convert='mp3_to_ogg')

# @application.route('/mp3towav', methods=['GET', 'POST'])
# def mp3_to_wav():
#     if request.method == 'POST':
#         try:
#             audio_file = request.files['file']
#             input_path = '/tmp/input.mp3'
#             output_path = '/tmp/output.wav'
#             audio_file.save(input_path)
#             ffmpeg.input(input_path).output(output_path).run()
#             return send_file(output_path, as_attachment=True, download_name='converted.wav')
#         except Exception as e:
#             print(f"Error: {e}")
#             abort(500)
#         finally:
#             if os.path.exists(input_path):
#                 os.remove(input_path)
#             if os.path.exists(output_path):
#                 os.remove(output_path)
#     return render_template('upload.html', convert='mp3_to_wav')


# Document converters
@application.route('/doctopdf', methods=['GET', 'POST'])
def doc_to_pdf():
    if request.method == 'POST':
        try:
            doc_file = request.files['file']
            temp_doc_path = 'temp.doc'
            temp_pdf_path = 'temp.pdf'
            doc_file.save(temp_doc_path)

            # Convert .doc to .pdf using LibreOffice
            convert_doc_to_pdf(temp_doc_path, temp_pdf_path)

            # Read the converted PDF file
            with open(temp_pdf_path, 'rb') as pdf_file:
                pdf_io = io.BytesIO(pdf_file.read())

            # Clean up temporary files
            os.remove(temp_doc_path)
            os.remove(temp_pdf_path)

            pdf_io.seek(0)
            return send_file(pdf_io, mimetype='application/pdf', as_attachment=True, download_name='converted.pdf')
        except Exception as e:
            print(f"Error: {e}")
            abort(500)
    return render_template('upload.html', convert='doc_to_pdf')
# Для pythonanywhere.com
# @application.route('/pdftodocx', methods=['GET', 'POST'])
# def pdf_to_docx():
#     if request.method == 'POST':
#         try:
#             pdf_file = request.files['file']
#             temp_pdf_path = '/tmp/temp.pdf'
#             output_path = '/tmp/output.docx'
            
#             # Save the uploaded PDF to a temporary location
#             pdf_file.save(temp_pdf_path)

#             # Read the PDF and convert to DOCX
#             pdf_reader = PdfReader(temp_pdf_path)
#             doc = Document()
#             for page in pdf_reader.pages:
#                 text = page.extract_text()
#                 if text:
#                     doc.add_paragraph(text)
            
#             # Save the DOCX file
#             doc.save(output_path)

#             # Send the DOCX file to the user
#             return send_file(output_path, as_attachment=True, download_name='converted.docx')
#         except Exception as e:
#             print(f"Error: {e}")
#             abort(500)
#     return render_template('upload.html', convert='pdf_to_docx')

@application.route('/pdftodocx', methods=['GET', 'POST'])
def pdf_to_docx():
    if request.method == 'POST':
        try:
            pdf_file = request.files['file']
            pdf_reader = PdfReader(pdf_file)
            doc = Document()
            for page in pdf_reader.pages:
                doc.add_paragraph(page.extract_text())
            output_path = 'output.docx'
            doc.save(output_path)
            return send_file(output_path, as_attachment=True, download_name='converted.docx')
        except Exception as e:
            print(f"Error: {e}")
            abort(500)
    return render_template('upload.html', convert='pdf_to_docx')

@application.route('/ppttopdf', methods=['GET', 'POST'])
def ppt_to_pdf():
    if request.method == 'POST':
        try:
            ppt_file = request.files['file']
            temp_ppt_path = 'temp.pptx'
            temp_pdf_path = 'temp.pdf'
            ppt_file.save(temp_ppt_path)

            # Convert .pptx to .pdf using LibreOffice
            convert_with_libreoffice(temp_ppt_path, 'pdf')

            with open(temp_pdf_path, 'rb') as pdf_file:
                pdf_io = io.BytesIO(pdf_file.read())

            os.remove(temp_ppt_path)
            os.remove(temp_pdf_path)

            pdf_io.seek(0)
            return send_file(pdf_io, mimetype='application/pdf', as_attachment=True, download_name='converted.pdf')
        except Exception as e:
            print(f"Error: {e}")
            abort(500)
    return render_template('upload.html', convert='ppt_to_pdf')

@application.route('/xlsxtopdf', methods=['GET', 'POST'])
def xlsx_to_pdf():
    if request.method == 'POST':
        try:
            xlsx_file = request.files['file']
            temp_xlsx_path = 'temp.xlsx'
            temp_pdf_path = 'temp.pdf'
            xlsx_file.save(temp_xlsx_path)

            # Convert .xlsx to .pdf using LibreOffice
            convert_with_libreoffice(temp_xlsx_path, 'pdf')

            with open(temp_pdf_path, 'rb') as pdf_file:
                pdf_io = io.BytesIO(pdf_file.read())

            os.remove(temp_xlsx_path)
            os.remove(temp_pdf_path)

            pdf_io.seek(0)
            return send_file(pdf_io, mimetype='application/pdf', as_attachment=True, download_name='converted.pdf')
        except Exception as e:
            print(f"Error: {e}")
            abort(500)
    return render_template('upload.html', convert='xlsx_to_pdf')

# Video converters
@application.route('/mp4toavi', methods=['GET', 'POST'])
def mp4_to_avi():
    if request.method == 'POST':
        try:
            video_file = request.files['file']
            input_path = 'input.mp4'
            output_path = 'output.avi'
            video_file.save(input_path)
            ffmpeg.input(input_path).output(output_path).run()
            return send_file(output_path, as_attachment=True, download_name='converted.avi')
        except Exception as e:
            print(f"Error: {e}")
            abort(500)
    return render_template('upload.html', convert='mp4_to_avi')

# Для pythonanywhere.com
# @application.route('/mp4toavi', methods=['GET', 'POST'])
# def mp4_to_avi():
#     if request.method == 'POST':
#         try:
#             video_file = request.files['file']
#             input_path = '/tmp/input.mp4'
#             output_path = '/tmp/output.avi'
#             video_file.save(input_path)

#             
#             ffmpeg.input(input_path).output(output_path).run()

#             # Отправка файла пользователю
#             return send_file(output_path, as_attachment=True, download_name='converted.avi')

#         except Exception as e:
#             print(f"Error: {e}")
#             abort(500)
#         finally:
#             # Удаление временных файлов
#             if os.path.exists(input_path):
#                 os.remove(input_path)
#             if os.path.exists(output_path):
#                 os.remove(output_path)

#     return render_template('upload.html', convert='mp4_to_avi')



@application.route('/avitomp4', methods=['GET', 'POST'])
def avi_to_mp4():
    if request.method == 'POST':
        try:
            video_file = request.files['file']
            input_path = 'input.avi'
            output_path = 'output.mp4'
            video_file.save(input_path)
            ffmpeg.input(input_path).output(output_path).run()
            return send_file(output_path, as_attachment=True, download_name='converted.mp4')
        except Exception as e:
            print(f"Error: {e}")
            abort(500)
    return render_template('upload.html', convert='avi_to_mp4')

@application.route('/movtomp4', methods=['GET', 'POST'])
def mov_to_mp4():
    if request.method == 'POST':
        try:
            video_file = request.files['file']
            input_bytes = io.BytesIO(video_file.read())
            output_bytes = io.BytesIO()

            # Use subprocess to call ffmpeg
            process = subprocess.Popen(
                ['ffmpeg', '-i', 'pipe:0', '-f', 'mp4', '-vcodec', 'libx264', '-preset', 'fast', '-movflags', 'frag_keyframe+empty_moov', 'pipe:1'],
                stdin=subprocess.PIPE,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE
            )

            output_data, error = process.communicate(input=input_bytes.getvalue())

            if process.returncode != 0:
                print(f"FFmpeg error: {error.decode()}")
                abort(500)

            output_bytes.write(output_data)
            output_bytes.seek(0)

            return send_file(output_bytes, mimetype='video/mp4', as_attachment=True, download_name='converted.mp4')
        except Exception as e:
            print(f"Error: {e}")
            abort(500)
    return render_template('upload.html', convert='mov_to_mp4')

@application.route('/mkvtomp4', methods=['GET', 'POST'])
def mkv_to_mp4():
    if request.method == 'POST':
        try:
            video_file = request.files['file']
            input_bytes = io.BytesIO(video_file.read())
            output_bytes = io.BytesIO()

            process = subprocess.Popen(
                ['ffmpeg', '-i', 'pipe:0', '-f', 'mp4', '-vcodec', 'libx264', '-preset', 'fast', '-movflags', 'frag_keyframe+empty_moov', 'pipe:1'],
                stdin=subprocess.PIPE,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE
            )

            output_data, error = process.communicate(input=input_bytes.getvalue())

            if process.returncode != 0:
                print(f"FFmpeg error: {error.decode()}")
                abort(500)

            output_bytes.write(output_data)
            output_bytes.seek(0)

            return send_file(output_bytes, mimetype='video/mp4', as_attachment=True, download_name='converted.mp4')
        except Exception as e:
            print(f"Error: {e}")
            abort(500)
    return render_template('upload.html', convert='mkv_to_mp4')

if __name__ == "__main__":
    application.run(host='0.0.0.0')


